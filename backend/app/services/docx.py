from __future__ import annotations

import io
from datetime import datetime

from docx import Document
from docx.shared import Pt

from app.schemas import Topic


def _add_heading(doc: Document, text: str, level: int) -> None:
    # python-docx supports Heading styles by name if available in the template.
    doc.add_heading(text, level=level)


def topics_to_docx_bytes(*, topics: list[Topic], source_label: str) -> bytes:
    doc = Document()

    # Title
    doc.add_paragraph("LexiNote")

    # Source + time
    doc.add_paragraph(f"Source: {source_label}")
    doc.add_paragraph(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}")
    doc.add_paragraph("")

    for topic in topics:
        _add_heading(doc, topic.title, level=2)
        for bullet in topic.bullets:
            # Bullet formatting: use list style.
            p = doc.add_paragraph(bullet.text, style="List Bullet")
            p.paragraph_format.space_after = Pt(0)

        doc.add_paragraph("")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

